from pathlib import Path
from typing import Union
import docx

def extract_text_from_docx(file_path: Union[str, Path]) -> str:
    """
    Extracts text from a DOCX file using python-docx.
    """
    path_str = str(file_path)
    try:
        doc = docx.Document(path_str)
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_str = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_str:
                    table_text.append(row_str)
        return "\n".join(paragraphs + table_text)
    except Exception:
        return ""
